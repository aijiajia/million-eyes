#!/usr/bin/env python
# encoding: utf-8


"""
@version: 
@author: Herbert
@license: Apache Licence 
@contact: 
@site: 
@software: PyCharm
@file: DocxDemo.py
@time: 2019/6/10 15:21
"""
import os
import re
from docx import Document
from docx.shared import Inches
from docx.dml.color import ColorFormat
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx():
    # 创建文档对象
    document = Document()

    # 设置默认字体
    document.styles['Normal'].font.name = '微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 创建自定义段落样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle1 = document.styles.add_style('UserStyle1', 1)
    # 设置字体尺寸
    UserStyle1.font.size = Pt(40)
    # 设置字体颜色
    UserStyle1.font.color.rgb = RGBColor(0xff, 0xde, 0x00)
    # 居中文本
    UserStyle1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置中文字体
    UserStyle1.font.name = '微软雅黑'
    UserStyle1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 创建自定义字符样式(第一个参数为样式名, 第二个参数为样式类型, 1为段落样式, 2为字符样式, 3为表格样式)
    UserStyle2 = document.styles.add_style('UserStyle2', 2)
    # 设置字体尺寸
    UserStyle2.font.size = Pt(15)
    # 设置字体颜色0c8ac5
    UserStyle2.font.color.rgb = RGBColor(0x0c, 0x8a, 0xc5)
    # 设置段落样式为宋体
    UserStyle2.font.name = '宋体'
    UserStyle2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 使用自定义段落样式
    document.add_paragraph('自定义段落样式', style=UserStyle1)

    # 使用自定义字符样式
    document.add_paragraph('').add_run('正月里采花无哟花采，二月间采花花哟正开，二月间采花花哟正开。三月里桃花红哟似海，四月间葡萄架哟上开，四月间葡萄架哟上开。', style=UserStyle2)

    # 设置粗体字
    document.add_paragraph('设置粗体字:').add_run('粗体字').bold = True

    # 设置斜体字
    document.add_paragraph('设置斜体字:').add_run('斜体字').italic = True

    # 设置字号50
    document.add_paragraph('设置字号50:').add_run('50').font.size = Pt(50)

    # 设置字体颜色为 af2626
    document.add_paragraph('设置字体颜色:').add_run('颜色').font.color.rgb = RGBColor(0xaf, 0x26, 0x26)

    # 样式叠加: 将字体改到30号并且将字体改成特定颜色;
    doubleStyle = document.add_paragraph('同时设置文字颜色和字号:').add_run('颜色和尺寸')
    doubleStyle.font.size = Pt(30)
    doubleStyle.font.color.rgb = RGBColor(0xaf, 0x26, 0x26)

    # 添加分页符
    document.add_page_break()

    # 创建 有序列表
    document.add_paragraph('').add_run('有序列表').font.size = Pt(30)
    document.add_paragraph('把冰箱门打开', style='List Number')
    document.add_paragraph('把大象装进去', style='List Number')
    document.add_paragraph('把冰箱门关上', style='List Number')

    # 创建 无序列表
    document.add_paragraph('').add_run('无序列表').font.size = Pt(30)
    document.add_paragraph('天地匆匆 惊鸿而过 路有千百个', style='List Bullet')
    document.add_paragraph('遑遑无归 闲云逸鹤 人间红尘过', style='List Bullet')
    document.add_paragraph('引势而流 鸿门乱局 各有各选择', style='List Bullet')
    document.add_paragraph('乾震坎艮 坤巽离兑 定一切生克', style='List Bullet')

    # 添加分页符
    document.add_page_break()
    # 添加图片
    document.add_paragraph('').add_run('添加图片').font.size = Pt(30)
    # document.add_picture('少女17087938.jpg', width=Inches(5))

    # 添加分页符
    document.add_page_break()

    document.add_paragraph('').add_run('创建表格').font.size = Pt(30)
    # 创建两行两列的表格
    rows_num = 5
    cols_num = 6
    table = document.add_table(rows=rows_num, cols=cols_num, style='Table Grid')

    for r in range(rows_num):
        for c in range(cols_num):
            table.cell(r, c).text = "第{r}行{c}列".format(r=r + 1, c=c + 1)
    # 保存文档
    document.save('Python生成的文档.docx')


def read_docx():
    # print(r'D:\work\长沙县应急项目\综合指挥调度系统解决方案V1-0-(2).docx')
    document = Document('Python修改的文档.docx')  # 打开文件demo.docx
    # for i in range(1, 200):
    #     print(i)
    #     print(document.paragraphs[i].text)  # 打印各段落内容文本
    # print(document.paragraphs[131].text)
    document.paragraphs[131].paragraph_format.line_spacing = 1.5
    # document.save('Python修改的文档.docx')
    print(document.paragraphs[130].paragraph_format.level)  # 打印各段落内容文本
    document.styles['Heading 1']


def para_select(docment):
    doc = Document(docment)
    print(doc.paragraphs[815].text)
    print(doc.paragraphs[816].text)
    # for i in range(1, len(doc.paragraphs)):
    #     if re.search(".*小处.*", doc.paragraphs[i].text):
    #         print(i)


def write_line(some_file, content):
    with open(some_file, 'a', encoding='utf-8') as f:
        f.write(content + '\n')


def head_content(document, some_file):
    doc = Document(document)
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1':
            write_line(some_file, 'Heading 1' + p.text)
        if p.style.name == 'Heading 2':
            write_line(some_file, 'Heading 2' + p.text)
        if p.style.name == 'Heading 3':
            write_line(some_file, 'Heading 3' + p.text)
        # if p.style.name == 'Heading 4':
        #     write_line(some_file, 'Heading 4' + p.text)
        # if p.style.name == 'Heading 5':
        #     write_line(some_file, 'Heading 5' + p.text)
        # if p.style.name == 'Heading 6':
        #     write_line(some_file, 'Heading 6' + p.text)


def extract_content_2_table(from_document):
    from_document = Document(from_document)
    to_document = Document()
    # 设置默认字体
    to_document.styles['Normal'].font.name = '微软雅黑'
    to_document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    to_document.add_paragraph('').add_run('创建表格').font.size = Pt(16)
    headings = [heading for heading in from_document.paragraphs if 'Heading' in heading.style.name]
    rows_num = len(headings)
    print(rows_num)
    cols_num = 6
    table = to_document.add_table(rows=rows_num, cols=cols_num, style='Table Grid')
    # table.font.name='宋体'
    # table.font.size = 140000
    for row in range(rows_num):
        print(row)
        # print(heading.text,int(heading.style.name[-1]))
        table.cell(row, int(headings[row].style.name[-1])-1).text = headings[row].text
    # 判断文件是否存在
    if (os.path.exists('生成表格.docx')):
        os.remove('生成表格.docx')
        print('文件已删除！')
    else:
        print("要删除的文件不存在！")
    to_document.save('生成表格.docx')


if __name__ == '__main__':
    # head_content(r'D:\work\甘肃省自然灾害监测预警体系与应急管理指挥系统项目\甘肃省应急管理综合应用平台建设方案20190702.docx')
    # para_select(r'C:\Users\herbert\Desktop\长沙县应急指挥平台体系建设方案cm0610.docx')
    # head_content(r'D:\work\科技查新20190603\4、参照样本：创我大数据分析服务平台—灵狐【编辑中】\【已有】12、需求说明书-创我大数据分析服务平台—灵狐项目.docx')
    # head_content(r'D:\work\甘肃省自然灾害监测预警体系与应急管理指挥系统项目\威而信-市应急管理局信息化建设方案（标准方案）(1).docx')
    head_content(r'D:\work\天水\天水智慧城管20190417\天水市2018智慧城市项目统筹建设方案V4.0--智慧城管分册.docx', 'somefile.txt')

    # extract_content_2_table(r'D:\work\天水\天水智慧城管20190417\天水市2018智慧城市项目统筹建设方案V4.0--智慧城管分册.docx')
